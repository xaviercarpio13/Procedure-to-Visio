from __future__ import annotations

from pathlib import Path
from docx import Document


def _clean(text: str) -> str:
    return " ".join((text or "").replace("\n", " ").split())


def extract_section6_tables(docx_path: str | Path) -> dict:
    """
    Extracts the procedure table from section 6:
    'CONTENIDO - DESCRIPCIÓN DEL PROCEDIMIENTO'.

    Returns:
        {
          "title": "...",
          "raw_text": "...",
          "rows": [{"no": "...", "actividad": "...", "descripcion": "...", "responsable": "..."}]
        }
    """
    docx_path = Path(docx_path)
    doc = Document(docx_path)

    title = docx_path.stem
    all_text = "\n".join(_clean(p.text) for p in doc.paragraphs if _clean(p.text))

    rows = []
    capture_started = False

    for table in doc.tables:
        table_rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
        flattened = " ".join(" ".join(r) for r in table_rows).upper()

        if "NO." in flattened and "ACTIVIDAD" in flattened and "RESPONSABLE" in flattened:
            capture_started = True

        if not capture_started:
            continue

        if not table_rows:
            continue

        header_idx = None
        for i, r in enumerate(table_rows):
            upper = [c.upper() for c in r]
            if any("NO" in c for c in upper) and any("ACTIVIDAD" in c for c in upper):
                header_idx = i
                break

        if header_idx is None:
            continue

        for r in table_rows[header_idx + 1:]:
            if len(r) < 4:
                continue

            no, actividad, descripcion, responsable = r[0], r[1], r[2], r[3]

            # Skip empty/invalid rows
            if not no and not actividad and not descripcion:
                continue
            if no.upper().startswith("NOMBRE") or "INDICADOR" in actividad.upper():
                continue

            rows.append({
                "no": no,
                "actividad": actividad,
                "descripcion": descripcion,
                "responsable": responsable,
            })

        # For this document structure, the section 6 table is one continuous table,
        # sometimes split across pages but still represented as one Word table.
        if rows:
            break

    if not rows:
        raise ValueError("Could not find the section 6 procedure table.")

    raw_text = "\n".join(
        f"{r['no']} | {r['actividad']} | {r['descripcion']} | {r['responsable']}"
        for r in rows
    )

    return {
        "title": title,
        "raw_text": raw_text,
        "rows": rows,
        "document_text": all_text,
    }
